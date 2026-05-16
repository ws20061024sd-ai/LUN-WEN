#!/usr/bin/env python3
"""
将正文.md 转换为排版规范的 Word 文档。
排版参照 CNKI 期刊论文样式。
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# ── 字体 ──
CN_BODY = "宋体"
CN_HEADING = "黑体"
CN_ABSTRACT = "楷体"
EN_FONT = "Times New Roman"

# ── 字号 ──
SIZE_TITLE = Pt(22)          # 二号
SIZE_HEADING = Pt(15)        # 小三
SIZE_BODY = Pt(12)           # 小四
SIZE_ABSTRACT = Pt(12)       # 小四
SIZE_NOTE = Pt(9)            # 小五
SIZE_AUTHOR = Pt(10.5)       # 五号

# ── 页面 ──
PAGE_W = Cm(21.0)
PAGE_H = Cm(29.7)
MARGIN_TOP = Cm(2.54)
MARGIN_BOTTOM = Cm(2.54)
MARGIN_LEFT = Cm(3.17)
MARGIN_RIGHT = Cm(3.17)

# ── 辅助函数 ──

def set_font(run, cn, en, size, bold=False):
    run.font.size = size
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), cn)
    rFonts.set(qn('w:ascii'), en)
    rFonts.set(qn('w:hAnsi'), en)
    rPr.insert(0, rFonts)

def p(doc, text="", cn=CN_BODY, en=EN_FONT, size=SIZE_BODY, bold=False,
      align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True, spacing=1.5,
      space_before=0, space_after=0, left_indent=None, right_indent=None):
    """添加段落。indent=None 表示不设首行缩进"""
    para = doc.add_paragraph()
    para.alignment = align
    pf = para.paragraph_format
    pf.line_spacing = spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if left_indent:
        pf.left_indent = left_indent
    if right_indent:
        pf.right_indent = right_indent
    if indent is True:
        pf.first_line_indent = Pt(24)   # 2 字符
    elif indent is False:
        pf.first_line_indent = Pt(0)
    if text:
        run = para.add_run(text)
        set_font(run, cn, en, size, bold)
    return para

def mixed_para(doc, segments, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True,
               spacing=1.5, space_before=0, space_after=0):
    """混合格式段落。segments = [(text, cn_font, en_font, size, bold), ...]"""
    para = doc.add_paragraph()
    para.alignment = align
    pf = para.paragraph_format
    pf.line_spacing = spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent is True:
        pf.first_line_indent = Pt(24)
    elif indent is False:
        pf.first_line_indent = Pt(0)
    for seg in segments:
        text, cn, en, size, bold = seg
        run = para.add_run(text)
        set_font(run, cn, en, size, bold)
    return para

def separator(doc):
    """注释前的分隔线"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = para.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.first_line_indent = Pt(0)
    run = para.add_run("—" * 20)
    set_font(run, CN_BODY, EN_FONT, Pt(9))

# ── 读取文件 ──
with open("/Users/xoln/Desktop/编程/论文/正文.md", "r") as f:
    lines = f.readlines()

# ── 创建文档 ──
doc = Document()

# 页面
for sec in doc.sections:
    sec.top_margin = MARGIN_TOP
    sec.bottom_margin = MARGIN_BOTTOM
    sec.left_margin = MARGIN_LEFT
    sec.right_margin = MARGIN_RIGHT
    sec.page_width = PAGE_W
    sec.page_height = PAGE_H

# ── 解析 ──
i = 0
abstract_lines = []
in_abstract = False
in_notes = False

while i < len(lines):
    line = lines[i].rstrip()

    if not line or line == "---":
        i += 1
        continue

    # ═══ 主标题 ═══
    if line.startswith("# ") and not line.startswith("## "):
        title = line[2:].strip()
        # 分隔行长度的装饰空行（留白）
        p(doc, "", size=Pt(8), align=WD_ALIGN_PARAGRAPH.CENTER,
          indent=False, space_before=0, space_after=0, spacing=1.0)
        p(doc, title, cn=CN_HEADING, en=EN_FONT, size=SIZE_TITLE, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, indent=False,
          space_before=0, space_after=24, spacing=1.2)
        # 作者占位
        p(doc, "作者姓名（单位，城市 邮编）", cn=CN_BODY, en=EN_FONT,
          size=SIZE_AUTHOR, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
          indent=False, space_before=0, space_after=18, spacing=1.2)
        i += 1
        continue

    # ═══ 二级标题 ═══
    if line.startswith("## "):
        heading_text = line[3:].strip()

        if heading_text == "摘要":
            in_abstract = True
            in_notes = False
            # 摘要标签
            p(doc, "摘要", cn=CN_HEADING, en=EN_FONT, size=SIZE_HEADING,
              bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False,
              space_before=6, space_after=8, spacing=1.5)
            i += 1
            continue

        if heading_text == "引言":
            in_abstract = False
            in_notes = False
            p(doc, "引言", cn=CN_HEADING, en=EN_FONT, size=SIZE_HEADING,
              bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False,
              space_before=12, space_after=6, spacing=1.5)
            i += 1
            continue

        if heading_text == "注释":
            in_abstract = False
            in_notes = True
            separator(doc)
            p(doc, "注释", cn=CN_HEADING, en=EN_FONT, size=SIZE_HEADING,
              bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False,
              space_before=6, space_after=8, spacing=1.5)
            i += 1
            continue

        # 一～六 一级标题
        if re.match(r'^[一二三四五六]、', heading_text):
            in_abstract = False
            in_notes = False
            p(doc, heading_text, cn=CN_HEADING, en=EN_FONT, size=SIZE_HEADING,
              bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False,
              space_before=14, space_after=6, spacing=1.5)
            i += 1
            continue

        # 其他二级标题
        in_abstract = False
        in_notes = False
        p(doc, heading_text, cn=CN_HEADING, en=EN_FONT, size=SIZE_HEADING,
          bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False,
          space_before=10, space_after=4, spacing=1.5)
        i += 1
        continue

    # ═══ 关键词 ═══
    if line.startswith("**关键词**"):
        kw = line.replace("**关键词**：", "").replace("**", "").strip()
        mixed_para(doc, [
            ("关键词：", CN_BODY, EN_FONT, SIZE_ABSTRACT, True),
            (kw, CN_BODY, EN_FONT, SIZE_ABSTRACT, False),
        ], indent=True, spacing=1.5, space_before=4, space_after=14)
        in_abstract = False
        i += 1
        continue

    # ═══ 正文 / 摘要 / 注释段落 ═══
    clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
    clean = re.sub(r'\^\((\d+)\)', r'(\1)', clean).strip()
    if not clean:
        i += 1
        continue

    if in_notes:
        p(doc, clean, cn=CN_BODY, en=EN_FONT, size=SIZE_NOTE,
          bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=False,
          spacing=1.15, space_before=0, space_after=1)
    elif in_abstract:
        # 摘要正文：左右缩进 + 楷体
        indent_em = Cm(1.0)
        p(doc, clean, cn=CN_ABSTRACT, en=EN_FONT, size=SIZE_ABSTRACT,
          bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True,
          spacing=1.5, space_before=0, space_after=0,
          left_indent=indent_em, right_indent=indent_em)
    else:
        p(doc, clean, cn=CN_BODY, en=EN_FONT, size=SIZE_BODY,
          bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True,
          spacing=1.5, space_before=0, space_after=0)

    i += 1

# ── 页码 ──
for sec in doc.sections:
    footer = sec.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run()
    f1 = OxmlElement('w:fldChar')
    f1.set(qn('w:fldCharType'), 'begin')
    r1._element.append(f1)
    r2 = fp.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    r2._element.append(instr)
    r3 = fp.add_run()
    f2 = OxmlElement('w:fldChar')
    f2.set(qn('w:fldCharType'), 'end')
    r3._element.append(f2)

# ── 输出 ──
out = "/Users/xoln/Desktop/编程/论文/正文.docx"
doc.save(out)
print(f"✅ 已生成: {out}")
